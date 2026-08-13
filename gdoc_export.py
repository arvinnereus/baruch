"""Export an AI Note to a .docx in the Google Drive sync folder.
Google Drive desktop syncs it up; the file opens directly in Google Docs.
No OAuth needed — pure local file write."""
import re
import time
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

def _drive_root() -> Path | None:
    """First Google Drive desktop mount found (any account)."""
    base = Path.home() / "Library/CloudStorage"
    if base.exists():
        for p in sorted(base.glob("GoogleDrive-*")):
            r = p / "My Drive"
            if r.exists():
                return r
    return None


def _export_dir(drive: Path | None) -> Path:
    """Drive/Baruch, but keep using an existing Drive/LocalFellow folder so the
    rename never orphans notes exported before it."""
    if not drive:
        return Path(__file__).resolve().parent / "data" / "exports"
    new, legacy = drive / "Baruch", drive / "LocalFellow"
    return legacy if (legacy.exists() and not new.exists()) else new


_drive = _drive_root()
DEFAULT_DIR = _export_dir(_drive)

TS_COLOR = RGBColor(0x4D, 0x6B, 0xFE)


def _add_rich(paragraph, text: str, ts: str | None = None):
    """Add text with **bold** spans, plus an optional timestamp chip."""
    for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1
    if ts:
        run = paragraph.add_run(f"  {ts}")
        run.font.color.rgb = TS_COLOR
        run.font.size = Pt(9)


def export_note(title: str, when: str, note: dict, out_dir: Path | None = None) -> Path:
    out_dir = out_dir or DEFAULT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(title, level=0)
    p = doc.add_paragraph()
    r = p.add_run(f"{when} · AI Notes by Baruch")
    r.italic = True
    r.font.size = Pt(9)

    doc.add_heading("Summary", level=1)
    for para in (note.get("summary") or "").split("\n\n"):
        if para.strip():
            _add_rich(doc.add_paragraph(), para.strip())

    if "action_items" in note:
        doc.add_heading("Action items", level=1)
        items = note.get("action_items") or []
        if items:
            for a in items:
                _add_rich(doc.add_paragraph(style="List Bullet"),
                          a.get("text", ""), a.get("ts"))
        else:
            doc.add_paragraph("No action items detected").runs[0].italic = True

    if "decisions" in note:
        doc.add_heading("Decisions", level=1)
        decs = note.get("decisions") or []
        if decs:
            for d in decs:
                _add_rich(doc.add_paragraph(style="List Bullet"),
                          d.get("text", ""), d.get("ts"))
        else:
            doc.add_paragraph("No decisions detected").runs[0].italic = True

    for sec in note.get("sections") or []:
        doc.add_heading(sec.get("title", "Section"), level=1)
        if sec.get("bullets"):
            for b in sec["bullets"]:
                _add_rich(doc.add_paragraph(style="List Bullet"),
                          b.get("text", ""), b.get("ts"))
        else:
            doc.add_paragraph("None noted").runs[0].italic = True

    if note.get("topics"):
        doc.add_heading("Topics", level=1)
        for t in note["topics"]:
            doc.add_heading(t.get("title", "Topic"), level=2)
            for b in t.get("bullets") or []:
                _add_rich(doc.add_paragraph(style="List Bullet"),
                          b.get("text", ""), b.get("ts"))

    safe = re.sub(r'[\\/:*?"<>|]', "-", title)[:80] or "Meeting"
    path = out_dir / f"{safe} — {time.strftime('%Y-%m-%d %H%M')}.docx"
    doc.save(path)
    return path
